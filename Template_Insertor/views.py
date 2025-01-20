from django.shortcuts import render
from Homepage.models import Headers,ProfessionalExperience,Education,Skills,Hobbies,Languages
from docx import Document
import os
from django.conf import settings
from datetime import date
import re

# Create your views here.

def callTemplateSelectorPage(request):
    current_user = request.user
    current_user_name = current_user.username

    Header_dict_user = Headers.objects.filter(user=current_user_name).values()[0]
    '''
    Headers.objects.filter(user=current_user_name) retrieves the fields and field values as a dictionary inside a list;
    Headers.objects.filter(user=current_user_name).values()[0] takes the details out of the list as a dictionary alone;
    we can now access the field names with Header_dict_user.keys() and the field values with Header_dict_user.values()

    we are doing the same for all the rest of the models too
    '''
    # # print(Header_of_user.values()[0])
    # print(Header_dict_user.keys())
    # print(Header_dict_user.values())

    ProfessionExperience_dict_user = ProfessionalExperience.objects.filter(user=current_user_name).values()[0]
    Education_dict_user = Education.objects.filter(user=current_user_name).values()[0]
    Skills_dict_user = Skills.objects.filter(user=current_user_name).values()[0]
    Languages_dict_user = Languages.objects.filter(user=current_user).values()[0]
    Hobbies_dict_user = Hobbies.objects.filter(user=current_user_name).values()[0]

    #merging all the dictionaries together
    merged_dict = {**Header_dict_user,**ProfessionExperience_dict_user,**Education_dict_user,**Skills_dict_user,**Languages_dict_user,**Hobbies_dict_user}
    # print(merged_dict)

    placeholder_dict = {}
    for key,value in merged_dict.items():
        if key not in ['id','created','user']:
            placeholder_dict[f"{{{{ {key} }}}}"] = value
    print(placeholder_dict)

    #starting with template modification

    template_path = os.path.join(settings.BASE_DIR,'Resume_Builder','resume_templates','template.docx')
    output_path = os.path.join(settings.BASE_DIR,'Resume_Builder','finished_Resumes','resume.docx')
    # print(template_path)
    # print(output_path)

    doc = Document(template_path)

    # Function to replace text in runs while maintaining formatting
    def replace_text_in_paragraph(paragraph, placeholder, replacement):
        full_text = ''.join(run.text for run in paragraph.runs)  # Combine all run texts
        if placeholder in full_text:

            #need to check if replacement is none. if replacement is none then we need to remove the entire placeholder corresponding to it in the template
            if replacement is None:
                #using regex to find the placeholder
                pattern = rf'\s*{re.escape(placeholder)}\s*' #matches placeholder and sorrunding white spaces
                full_text = re.sub(pattern,'',full_text)

            else:
                #checking if the replacement is a date so that it needs to be converted in to a string
                if isinstance(replacement,date):
                    replacement = replacement.strftime('%d-%m-%Y')


                full_text = full_text.replace(placeholder, replacement)  # Replace placeholder
            for run in paragraph.runs:
                run.text = ''  # Clear current runs
            paragraph.runs[0].text = full_text  # Apply the updated text to the first run

    # Iterate through paragraphs and replace placeholders
    for paragraph in doc.paragraphs:
        for placeholder, replacement in placeholder_dict.items():
            replace_text_in_paragraph(paragraph, placeholder, replacement)

    # Iterate through tables and replace placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, replacement in placeholder_dict.items():
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, placeholder, replacement)
    

    doc.save(output_path)
    print('resume created succesfully')

    return render(request,'Template_Insertor/template_selector.html',{})
